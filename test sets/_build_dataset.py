"""Temporary script to build dataset_json from Test_set_01 docx files."""
import os, re, json
from pathlib import Path
from docx import Document

TEST_SET_DIR = 'Test_set_01'
DATASET_DIR = 'dataset_json'
IMAGE_DIR = 'images'

os.makedirs(DATASET_DIR, exist_ok=True)
os.makedirs(IMAGE_DIR, exist_ok=True)

# FIX: lines at or after these markers are AAMC footer metadata, not question content
FOOTER_MARKERS = [
    'Khan Academy Lessons',
    'Content & Skills',
    'Content &\xa0Skills',
    'Provide Feedback on AAMC',
    '© 2026 AAMC',
    'Contact Us Web Accessibility',
]

# FIX: lines matching these patterns are AAMC navigation headers, strip them
HEADER_MARKERS = [
    'Skip to main content',
    'Unscored Sample Test',
    'Done Reviewing',
]


def extract_text(docx_path):
    doc = Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return '\n'.join(paragraphs)


def clean_text(text):
    """FIX: strip AAMC header/footer metadata that confuses the parser."""
    lines = text.split('\n')

    # Strip header lines
    cleaned = []
    past_header = False
    for line in lines:
        if not past_header and any(line.startswith(m) for m in HEADER_MARKERS):
            continue
        past_header = True
        # Stop at footer markers
        if any(m in line for m in FOOTER_MARKERS):
            break
        cleaned.append(line)

    return '\n'.join(cleaned)


def extract_images(docx_path, question_id):
    doc = Document(docx_path)
    img_paths = []
    for rel in doc.part.rels.values():
        if 'image' in rel.target_ref:
            img_data = rel.target_part.blob
            img_name = f'{question_id}_{len(img_paths)}.png'
            img_path = os.path.join(IMAGE_DIR, img_name)
            with open(img_path, 'wb') as f:
                f.write(img_data)
            img_paths.append(img_path)
    return img_paths


def parse_question(text):
    # FIX: strip AAMC header/footer metadata before parsing
    text = clean_text(text)

    # Extract answer first (before we lose the Solution line)
    answer_match = re.search(r'correct answer is\s*([A-D])', text, re.I)
    answer = answer_match.group(1).upper() if answer_match else None

    # FIX: truncate at "Solution:" line — everything after is explanation, not question
    solution_match = re.search(r'\nSolution:', text)
    if solution_match:
        text = text[:solution_match.start()]

    # FIX: handle multi-line choice format where "A." is on one line and choice text on next
    # Pattern 1: "A.\n<choice text>" (AAMC format — letter alone on line, text on next)
    # Pattern 2: "A. <choice text>" (inline format)
    lines = text.split('\n')
    choices = {}
    choice_lines_start = None
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        # Check for "A." alone on a line
        m = re.match(r'^([A-D])\.\s*$', line)
        if m:
            letter = m.group(1)
            # FIX: next line might be another choice letter (image-only) or choice text
            if i + 1 < len(lines) and not re.match(r'^[A-D]\.', lines[i + 1].strip()):
                choice_text = lines[i + 1].strip()
                choices[letter] = choice_text
                if choice_lines_start is None:
                    choice_lines_start = i
                i += 2
            else:
                # FIX: image-only choice — no text, choice is a diagram
                choices[letter] = '[image]'
                if choice_lines_start is None:
                    choice_lines_start = i
                i += 1
            continue
        # FIX: Check for "A. <text>" or "A.<text>" inline (space optional)
        m = re.match(r'^([A-D])\.\s*(\S.+)', line)
        if m:
            letter = m.group(1)
            choice_text = m.group(2).strip()
            choices[letter] = choice_text
            if choice_lines_start is None:
                choice_lines_start = i
            i += 1
            continue
        i += 1

    # Everything before the choices block is passage + question
    if choice_lines_start is not None:
        pre_choices = '\n'.join(lines[:choice_lines_start])
    else:
        pre_choices = text

    # FIX: find the question — last line ending with '?' or ':' before choices
    question = ''
    pre_lines = pre_choices.strip().split('\n')
    for j in range(len(pre_lines) - 1, -1, -1):
        line = pre_lines[j].strip()
        if line.endswith('?') or line.endswith(':'):
            question = line
            passage = '\n'.join(pre_lines[:j]).strip()
            break
    else:
        # No question line found — use the last non-empty line as question
        passage = pre_choices.strip()

    return passage, question, choices, answer

# Collect docx files
docx_files = []
for section in sorted(os.listdir(TEST_SET_DIR)):
    section_path = os.path.join(TEST_SET_DIR, section)
    if os.path.isdir(section_path):
        for f in sorted(os.listdir(section_path)):
            if f.lower().endswith('.docx') and not f.startswith('~$'):
                docx_files.append(os.path.join(section_path, f))

print(f'Found {len(docx_files)} DOCX files in {TEST_SET_DIR}')
for section in sorted(os.listdir(TEST_SET_DIR)):
    section_path = os.path.join(TEST_SET_DIR, section)
    if os.path.isdir(section_path):
        count = len([f for f in os.listdir(section_path) if f.lower().endswith('.docx')])
        print(f'  {section}: {count} questions')

# Build dataset
dataset = []
for i, file in enumerate(docx_files):
    question_id = f'q{i}'
    text = extract_text(file)
    images = extract_images(file, question_id)
    passage, question, choices, answer = parse_question(text)
    source_path = Path(file)
    sample = {
        'id': question_id,
        'passage': passage,
        'question': question,
        'choices': choices,
        'images': images,
        'answer': answer,
        'source_file': source_path.name,
        'source_stem': source_path.stem,
        'test_set': 'test_set_01',
        'subject_category': source_path.parent.name,
    }
    dataset.append(sample)
    with open(os.path.join(DATASET_DIR, f'{question_id}.json'), 'w', encoding='utf-8') as f:
        json.dump(sample, f, indent=2)

print(f'\nDataset built with {len(dataset)} samples')
n_no_answer = sum(1 for s in dataset if not s.get('answer'))
n_no_question = sum(1 for s in dataset if not s.get('question'))
n_no_choices = sum(1 for s in dataset if len(s.get('choices', {})) < 4)
n_with_images = sum(1 for s in dataset if s.get('images'))
print(f'Validation: {n_no_answer} missing answers, {n_no_question} missing questions, {n_no_choices} incomplete choice sets, {n_with_images} with images')

# Show first 3 samples
for s in dataset[:3]:
    print(f"\n--- {s['id']} ({s['source_file']}, {s['subject_category']}) ---")
    print(f"  Q: {s['question'][:120]}")
    print(f"  Choices: {list(s['choices'].keys())}")
    print(f"  Answer: {s['answer']}")
    print(f"  Images: {len(s['images'])}")
