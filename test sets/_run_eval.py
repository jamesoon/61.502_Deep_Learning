"""
Full evaluation runner for MCAT benchmark — test_set_01 with Qwen3-14B.
Executes all notebook cells in sequence as a script.
"""
import os
import re
import json
import glob
import time
from pathlib import Path
from typing import Dict, List, Optional

import numpy as np
import pandas as pd
from openai import OpenAI
from docx import Document

print("=" * 60)
print("MCAT Local Model Benchmark — test_set_01")
print("=" * 60)

# ======================================================================
# CELL 1: Dataset building
# ======================================================================

TEST_SET_DIR = 'Test_set_01'
EVAL_SPLIT_NAME_BUILD = 'test_set_01'
DATASET_DIR = 'dataset_json'
IMAGE_DIR = 'images'

os.makedirs(DATASET_DIR, exist_ok=True)
os.makedirs(IMAGE_DIR, exist_ok=True)

FOOTER_MARKERS = [
    'Khan Academy Lessons',
    'Content & Skills',
    'Content &\xa0Skills',
    'Provide Feedback on AAMC',
    '\u00a9 2026 AAMC',
    'Contact Us Web Accessibility',
]
HEADER_MARKERS = [
    'Skip to main content',
    'Unscored Sample Test',
    'Done Reviewing',
]


def extract_text(docx_path):
    doc = Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return '\n'.join(paragraphs)


def clean_text(text):
    lines = text.split('\n')
    cleaned = []
    past_header = False
    for line in lines:
        if not past_header and any(line.startswith(m) for m in HEADER_MARKERS):
            continue
        past_header = True
        if any(m in line for m in FOOTER_MARKERS):
            break
        cleaned.append(line)
    return '\n'.join(cleaned)


def extract_images(docx_path, question_id):
    doc = Document(docx_path)
    img_paths = []
    for rel in doc.part.rels.values():
        if 'image' in rel.target_ref:
            img_data = rel.target_part.blob
            img_name = f'{question_id}_{len(img_paths)}.png'
            img_path = os.path.join(IMAGE_DIR, img_name)
            with open(img_path, 'wb') as f:
                f.write(img_data)
            img_paths.append(img_path)
    return img_paths


def parse_question(text):
    text = clean_text(text)
    answer_match = re.search(r'correct answer is\s*([A-D])', text, re.I)
    answer = answer_match.group(1).upper() if answer_match else None
    solution_match = re.search(r'\nSolution:', text)
    if solution_match:
        text = text[:solution_match.start()]
    lines = text.split('\n')
    choices = {}
    choice_lines_start = None
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        m = re.match(r'^([A-D])\.\s*$', line)
        if m:
            letter = m.group(1)
            if i + 1 < len(lines) and not re.match(r'^[A-D]\.', lines[i + 1].strip()):
                choice_text = lines[i + 1].strip()
                choices[letter] = choice_text
                if choice_lines_start is None:
                    choice_lines_start = i
                i += 2
            else:
                choices[letter] = '[image]'
                if choice_lines_start is None:
                    choice_lines_start = i
                i += 1
            continue
        m = re.match(r'^([A-D])\.\s*(\S.+)', line)
        if m:
            letter = m.group(1)
            choice_text = m.group(2).strip()
            choices[letter] = choice_text
            if choice_lines_start is None:
                choice_lines_start = i
            i += 1
            continue
        i += 1
    if choice_lines_start is not None:
        pre_choices = '\n'.join(lines[:choice_lines_start])
    else:
        pre_choices = text
    question = ''
    pre_lines = pre_choices.strip().split('\n')
    for j in range(len(pre_lines) - 1, -1, -1):
        line = pre_lines[j].strip()
        if line.endswith('?') or line.endswith(':'):
            question = line
            passage = '\n'.join(pre_lines[:j]).strip()
            break
    else:
        passage = pre_choices.strip()
    return passage, question, choices, answer


def build_dataset(docx_files, test_set_name=None):
    dataset = []
    for i, file in enumerate(docx_files):
        question_id = f'q{i}'
        text = extract_text(file)
        images = extract_images(file, question_id)
        passage, question, choices, answer = parse_question(text)
        source_path = Path(file)
        sample = {
            'id': question_id,
            'passage': passage,
            'question': question,
            'choices': choices,
            'images': images,
            'answer': answer,
            'source_file': source_path.name,
            'source_stem': source_path.stem,
            'test_set': test_set_name or 'default_test_set',
            'subject_category': source_path.parent.name,
        }
        dataset.append(sample)
        with open(os.path.join(DATASET_DIR, f'{question_id}.json'), 'w', encoding='utf-8') as f:
            json.dump(sample, f, indent=2)
    return dataset


# Collect and build
docx_files = []
for section in sorted(os.listdir(TEST_SET_DIR)):
    section_path = os.path.join(TEST_SET_DIR, section)
    if os.path.isdir(section_path):
        for f in sorted(os.listdir(section_path)):
            if f.lower().endswith('.docx') and not f.startswith('~$'):
                docx_files.append(os.path.join(section_path, f))

print(f'\n[STEP 1] Found {len(docx_files)} DOCX files in {TEST_SET_DIR}')
for section in sorted(os.listdir(TEST_SET_DIR)):
    section_path = os.path.join(TEST_SET_DIR, section)
    if os.path.isdir(section_path):
        count = len([f for f in os.listdir(section_path) if f.lower().endswith('.docx') and not f.startswith('~$')])
        print(f'  {section}: {count} questions')

built_dataset = build_dataset(docx_files, test_set_name=EVAL_SPLIT_NAME_BUILD)
print(f'Dataset built with {len(built_dataset)} samples')
n_no_answer = sum(1 for s in built_dataset if not s.get('answer'))
n_no_choices = sum(1 for s in built_dataset if len(s.get('choices', {})) < 4)
print(f'Validation: {n_no_answer} missing answers, {n_no_choices} incomplete choice sets')

# ======================================================================
# CELL 3: Evaluation config + functions
# ======================================================================

LOCAL_API_BASE_URL = os.getenv("LOCAL_API_BASE_URL", "http://127.0.0.1:1234/v1")
LOCAL_API_KEY = os.getenv("LOCAL_API_KEY", "not-needed")
client = OpenAI(base_url=LOCAL_API_BASE_URL, api_key=LOCAL_API_KEY)

RESULT_DIR = "results_local_llm"
PER_MODEL_DIR = os.path.join(RESULT_DIR, "per_model")
FIGURE_DIR = os.path.join(RESULT_DIR, "figures")

PREDICTIONS_PATH = os.path.join(RESULT_DIR, "all_predictions.csv")
SUMMARY_BY_SPLIT_PATH = os.path.join(RESULT_DIR, "summary_by_split.csv")
SUMMARY_AGG_PATH = os.path.join(RESULT_DIR, "summary_aggregated.csv")
SUBJECT_SUMMARY_PATH = os.path.join(RESULT_DIR, "summary_by_subject.csv")

os.makedirs(RESULT_DIR, exist_ok=True)
os.makedirs(PER_MODEL_DIR, exist_ok=True)
os.makedirs(FIGURE_DIR, exist_ok=True)

EVAL_SPLIT_NAME = "test_set_01"
EXPERIMENT_TAG = "mcat_local_benchmark_v2"
TEMPERATURE = 0.0
# FIX: Qwen3 uses <think> reasoning by default — need enough tokens for think + answer
MAX_TOKENS = 2048
SLEEP_BETWEEN_REQUESTS = 0.0
LABELS = ["A", "B", "C", "D"]

MODEL_SPECS = {
    "qwen3_14b": {
        "display_name": "Qwen3-14B-Q4_K_M",
        "hf_repo": "Qwen/Qwen3-14B-GGUF",
        "hf_file": "Qwen3-14B-Q4_K_M.gguf",
        "match_tokens": ["qwen3", "14b"],
    },
    "gemma3_12b": {
        "display_name": "Gemma-3-12B-it-q4_0",
        "hf_repo": "google/gemma-3-12b-it-qat-q4_0-gguf",
        "hf_file": "gemma-3-12b-it-q4_0.gguf",
        "match_tokens": ["gemma", "3", "12b"],
    },
    "gpt_oss_20b": {
        "display_name": "gpt-oss-20b-MXFP4",
        "hf_repo": "bartowski/openai_gpt-oss-20b-GGUF",
        "hf_file": "gpt-oss-20b-MXFP4.gguf",
        "match_tokens": ["gpt", "oss", "20b"],
    },
}

SYSTEM_PROMPT = (
    "You are taking an MCAT multiple-choice benchmark. "
    "Solve the question carefully, but return only one capital letter: A, B, C, or D. "
    "Do not explain your reasoning."
)


def normalize_text(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", str(text).lower())


def list_server_models() -> List[str]:
    try:
        response = client.models.list()
        data = getattr(response, "data", []) or []
        return [getattr(item, "id", str(item)) for item in data]
    except Exception as e:
        print(f"Could not list models from local server: {e}")
        return []


def resolve_loaded_model_id(spec, available_model_ids):
    if not available_model_ids:
        return None
    normalized_tokens = [normalize_text(token) for token in spec["match_tokens"]]
    for model_id in available_model_ids:
        normalized_id = normalize_text(model_id)
        if all(token in normalized_id for token in normalized_tokens):
            return model_id
    return None


def infer_subject_category(sample):
    for key in ["subject_category", "subject", "category", "section", "topic"]:
        value = sample.get(key)
        if value is not None and str(value).strip():
            return str(value).strip()
    return None


def load_dataset(dataset_dir=DATASET_DIR, eval_split_name=EVAL_SPLIT_NAME):
    dataset = []
    for path in sorted(glob.glob(os.path.join(dataset_dir, "*.json"))):
        with open(path, "r", encoding="utf-8") as f:
            sample = json.load(f)
        sample["source_json"] = os.path.basename(path)
        sample["eval_split_name"] = sample.get("test_set") or eval_split_name
        sample["subject_category"] = infer_subject_category(sample)
        dataset.append(sample)
    return dataset


def build_question_prompt(sample):
    choices = sample.get("choices", {})
    return f"""Passage:
{sample.get('passage', '').strip()}

Question:
{sample.get('question', '').strip()}

Choices:
A. {choices.get('A', '')}
B. {choices.get('B', '')}
C. {choices.get('C', '')}
D. {choices.get('D', '')}

Return only one capital letter: A, B, C, or D.
"""


def extract_letter(raw_text):
    if raw_text is None:
        return None
    text = str(raw_text).strip()
    # FIX: Qwen3 wraps reasoning in <think>...</think> — strip it to get the actual answer
    think_end = text.rfind('</think>')
    if think_end != -1:
        text = text[think_end + len('</think>'):].strip()
    text = text.upper()
    if text in {"A", "B", "C", "D"}:
        return text
    match = re.search(r"\b([A-D])\b", text)
    if match:
        return match.group(1)
    match = re.search(r"ANSWER\s*[:\-]?\s*([A-D])", text)
    if match:
        return match.group(1)
    return None


def query_local_model(server_model_id, sample, temperature=TEMPERATURE, max_tokens=MAX_TOKENS):
    response = client.chat.completions.create(
        model=server_model_id,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": build_question_prompt(sample)},
        ],
        temperature=temperature,
        max_tokens=max_tokens,
    )
    raw_text = response.choices[0].message.content
    pred = extract_letter(raw_text)
    return pred, raw_text


def compute_macro_classification_metrics(df):
    y_true = df["gold_answer"].fillna("INVALID")
    y_pred = df["pred_answer"].fillna("INVALID")
    precisions, recalls, f1s = [], [], []
    for label in LABELS:
        tp = int(((y_true == label) & (y_pred == label)).sum())
        fp = int(((y_true != label) & (y_pred == label)).sum())
        fn = int(((y_true == label) & (y_pred != label)).sum())
        precision = tp / (tp + fp) if (tp + fp) > 0 else 0.0
        recall = tp / (tp + fn) if (tp + fn) > 0 else 0.0
        f1 = (2 * precision * recall / (precision + recall)) if (precision + recall) > 0 else 0.0
        precisions.append(precision)
        recalls.append(recall)
        f1s.append(f1)
    return {
        "macro_precision": float(np.mean(precisions)),
        "macro_recall": float(np.mean(recalls)),
        "macro_f1": float(np.mean(f1s)),
    }


def evaluate_model(model_key, server_model_id, dataset, sleep_between_requests=SLEEP_BETWEEN_REQUESTS):
    rows = []
    print(f'\n[EVAL] Evaluating {model_key} on split={EVAL_SPLIT_NAME} using server id: {server_model_id}')
    for idx, sample in enumerate(dataset, start=1):
        t0 = time.perf_counter()
        pred, raw_text = query_local_model(server_model_id, sample)
        latency_s = time.perf_counter() - t0
        gold = (sample.get("answer") or "").strip().upper() or None
        correct = bool(pred == gold) if (pred is not None and gold is not None) else False
        row = {
            "experiment_tag": EXPERIMENT_TAG,
            "eval_split_name": sample.get("eval_split_name", EVAL_SPLIT_NAME),
            "model_key": model_key,
            "model_display_name": MODEL_SPECS[model_key]["display_name"],
            "server_model_id": server_model_id,
            "hf_repo": MODEL_SPECS[model_key]["hf_repo"],
            "hf_file": MODEL_SPECS[model_key]["hf_file"],
            "temperature": TEMPERATURE,
            "max_tokens": MAX_TOKENS,
            "question_id": sample.get("id"),
            "source_json": sample.get("source_json"),
            "source_file": sample.get("source_file"),
            "subject_category": sample.get("subject_category"),
            "gold_answer": gold,
            "pred_answer": pred,
            "correct": correct,
            "latency_s": round(latency_s, 4),
            "raw_output": raw_text,
            "question": sample.get("question", ""),
        }
        rows.append(row)
        print(
            f'  [{idx}/{len(dataset)}] {sample.get("id")} | '
            f'gold={gold} | pred={pred} | {"OK" if correct else "WRONG"} | {latency_s:.2f}s'
        )
        if sleep_between_requests > 0:
            time.sleep(sleep_between_requests)
    df = pd.DataFrame(rows)
    per_model_path = os.path.join(PER_MODEL_DIR, f'{EVAL_SPLIT_NAME}__{model_key}.csv')
    df.to_csv(per_model_path, index=False)
    print(f'  Saved per-question results to: {per_model_path}')
    return df


def upsert_predictions(run_df, predictions_path=PREDICTIONS_PATH):
    if run_df.empty:
        return run_df.copy()
    key_cols = ["experiment_tag", "eval_split_name", "model_key", "question_id"]
    if os.path.exists(predictions_path):
        existing = pd.read_csv(predictions_path)
        merged = pd.concat([existing, run_df], ignore_index=True)
    else:
        merged = run_df.copy()
    merged = merged.sort_values(by=key_cols).drop_duplicates(subset=key_cols, keep="last").reset_index(drop=True)
    merged.to_csv(predictions_path, index=False)
    return merged


def build_summary_by_split(pred_df):
    if pred_df.empty:
        return pd.DataFrame()
    rows = []
    group_cols = [
        "experiment_tag", "eval_split_name", "model_key", "model_display_name",
        "server_model_id", "hf_repo", "hf_file", "temperature", "max_tokens"
    ]
    for keys, grp in pred_df.groupby(group_cols, dropna=False):
        metrics = compute_macro_classification_metrics(grp)
        row = dict(zip(group_cols, keys))
        row.update({
            "n_questions": int(len(grp)),
            "n_correct": int(grp["correct"].sum()),
            "accuracy": float(grp["correct"].mean()),
            "invalid_prediction_rate": float((~grp["pred_answer"].isin(LABELS)).mean()),
            "mean_latency_s": float(grp["latency_s"].mean()),
            "median_latency_s": float(grp["latency_s"].median()),
            "p90_latency_s": float(grp["latency_s"].quantile(0.90)),
        })
        row.update(metrics)
        rows.append(row)
    summary = pd.DataFrame(rows).sort_values(
        by=["eval_split_name", "accuracy", "macro_f1", "mean_latency_s"],
        ascending=[True, False, False, True]
    ).reset_index(drop=True)
    for col in ["accuracy", "macro_precision", "macro_recall", "macro_f1", "invalid_prediction_rate", "mean_latency_s", "median_latency_s", "p90_latency_s"]:
        summary[col] = summary[col].round(4)
    return summary


def build_summary_aggregated(summary_by_split_df):
    if summary_by_split_df.empty:
        return pd.DataFrame()
    agg = summary_by_split_df.groupby(
        ["experiment_tag", "model_key", "model_display_name", "hf_repo", "hf_file"], as_index=False
    ).agg(
        n_test_sets=("eval_split_name", "nunique"),
        total_questions=("n_questions", "sum"),
        mean_accuracy=("accuracy", "mean"),
        std_accuracy=("accuracy", "std"),
        mean_macro_precision=("macro_precision", "mean"),
        mean_macro_recall=("macro_recall", "mean"),
        mean_macro_f1=("macro_f1", "mean"),
        mean_invalid_prediction_rate=("invalid_prediction_rate", "mean"),
        mean_latency_s=("mean_latency_s", "mean"),
        median_latency_s=("median_latency_s", "mean"),
        mean_p90_latency_s=("p90_latency_s", "mean"),
    )
    agg["std_accuracy"] = agg["std_accuracy"].fillna(0.0)
    for col in ["mean_accuracy", "std_accuracy", "mean_macro_precision", "mean_macro_recall", "mean_macro_f1", "mean_invalid_prediction_rate", "mean_latency_s", "median_latency_s", "mean_p90_latency_s"]:
        agg[col] = agg[col].round(4)
    return agg.sort_values(by=["mean_accuracy", "mean_macro_f1", "mean_latency_s"], ascending=[False, False, True]).reset_index(drop=True)


def build_subject_summary(pred_df):
    if pred_df.empty or "subject_category" not in pred_df.columns:
        return pd.DataFrame()
    subject_df = pred_df.copy()
    subject_df["subject_category"] = subject_df["subject_category"].fillna("Unknown")
    rows = []
    group_cols = ["experiment_tag", "eval_split_name", "model_key", "model_display_name", "subject_category"]
    for keys, grp in subject_df.groupby(group_cols, dropna=False):
        metrics = compute_macro_classification_metrics(grp)
        row = dict(zip(group_cols, keys))
        row.update({
            "n_questions": int(len(grp)),
            "accuracy": float(grp["correct"].mean()),
            "mean_latency_s": float(grp["latency_s"].mean()),
        })
        row.update(metrics)
        rows.append(row)
    subject_summary = pd.DataFrame(rows).sort_values(by=["subject_category", "accuracy"], ascending=[True, False]).reset_index(drop=True)
    for col in ["accuracy", "mean_latency_s", "macro_precision", "macro_recall", "macro_f1"]:
        subject_summary[col] = subject_summary[col].round(4)
    return subject_summary


# ======================================================================
# CELL 4: Model discovery
# ======================================================================

print(f'\n[STEP 2] Discovering models on {LOCAL_API_BASE_URL}')
available_model_ids = list_server_models()
print(f'  Available: {available_model_ids}')

resolved_models = {}
for model_key, spec in MODEL_SPECS.items():
    resolved_id = resolve_loaded_model_id(spec, available_model_ids)
    if resolved_id is not None:
        resolved_models[model_key] = resolved_id

print(f'  Resolved targets: {resolved_models}')
if not resolved_models:
    raise RuntimeError('No target models matched. Load a model in LM Studio first.')

# ======================================================================
# CELL 6: Evaluation run
# ======================================================================

print(f'\n[STEP 3] Loading dataset and running evaluation')
eval_dataset = load_dataset(DATASET_DIR, EVAL_SPLIT_NAME)
print(f'  Loaded {len(eval_dataset)} questions for split={EVAL_SPLIT_NAME}')

if len(eval_dataset) == 0:
    raise RuntimeError('No dataset_json/*.json files found.')

all_run_dfs = []
for model_key, server_model_id in resolved_models.items():
    run_df = evaluate_model(model_key, server_model_id, eval_dataset)
    all_run_dfs.append(run_df)

if all_run_dfs:
    current_run_df = pd.concat(all_run_dfs, ignore_index=True)
    all_predictions_df = upsert_predictions(current_run_df, PREDICTIONS_PATH)

    summary_by_split_df = build_summary_by_split(all_predictions_df)
    summary_aggregated_df = build_summary_aggregated(summary_by_split_df)
    subject_summary_df = build_subject_summary(all_predictions_df)

    summary_by_split_df.to_csv(SUMMARY_BY_SPLIT_PATH, index=False)
    summary_aggregated_df.to_csv(SUMMARY_AGG_PATH, index=False)
    subject_summary_df.to_csv(SUBJECT_SUMMARY_PATH, index=False)

    print('\n[STEP 4] Summary by test set:')
    print(summary_by_split_df.to_string(index=False))

    print('\n[STEP 5] Aggregated summary:')
    print(summary_aggregated_df.to_string(index=False))

    if not subject_summary_df.empty:
        print('\n[STEP 6] Subject/category summary:')
        print(subject_summary_df.to_string(index=False))

# ======================================================================
# CELL 8: Visualization
# ======================================================================

import matplotlib
matplotlib.use('Agg')  # Non-interactive backend for script execution
import matplotlib.pyplot as plt


def ensure_nonempty_csv(path):
    if not os.path.exists(path):
        print(f'  File not found: {path}')
        return pd.DataFrame()
    df = pd.read_csv(path)
    if df.empty:
        print(f'  File is empty: {path}')
    return df


def save_grouped_bar(df, x_col, y_cols, title, ylabel, save_path):
    if df.empty:
        return
    x_labels = df[x_col].astype(str).tolist()
    x = np.arange(len(x_labels))
    width = 0.8 / max(len(y_cols), 1)
    plt.figure(figsize=(10, 5))
    for i, y_col in enumerate(y_cols):
        plt.bar(x + i * width, df[y_col], width=width, label=y_col)
    plt.xticks(x + width * (len(y_cols) - 1) / 2, x_labels, rotation=15)
    plt.ylabel(ylabel)
    plt.title(title)
    plt.legend()
    plt.tight_layout()
    plt.savefig(save_path, dpi=200, bbox_inches="tight")
    plt.close()
    print(f'  Saved: {save_path}')


def save_latency_boxplot(pred_df, save_path):
    if pred_df.empty:
        return
    model_order = pred_df.groupby("model_display_name")["latency_s"].mean().sort_values().index.tolist()
    data = [pred_df.loc[pred_df["model_display_name"] == mn, "latency_s"].dropna().values for mn in model_order]
    plt.figure(figsize=(10, 5))
    plt.boxplot(data, labels=model_order, showfliers=False)
    plt.ylabel("Latency (seconds)")
    plt.title("Latency distribution by model")
    plt.xticks(rotation=15)
    plt.tight_layout()
    plt.savefig(save_path, dpi=200, bbox_inches="tight")
    plt.close()
    print(f'  Saved: {save_path}')


def save_accuracy_by_test_set(summary_by_split_df, save_path):
    if summary_by_split_df.empty:
        return
    pivot = summary_by_split_df.pivot_table(index="eval_split_name", columns="model_display_name", values="accuracy", aggfunc="mean").sort_index()
    if pivot.empty:
        return
    x_labels = pivot.index.astype(str).tolist()
    model_names = pivot.columns.tolist()
    x = np.arange(len(x_labels))
    width = 0.8 / max(len(model_names), 1)
    plt.figure(figsize=(11, 5))
    for i, mn in enumerate(model_names):
        plt.bar(x + i * width, pivot[mn].fillna(0.0).values, width=width, label=mn)
    plt.xticks(x + width * (len(model_names) - 1) / 2, x_labels, rotation=15)
    plt.ylim(0, 1)
    plt.ylabel("Accuracy")
    plt.title("Accuracy by test set")
    plt.legend()
    plt.tight_layout()
    plt.savefig(save_path, dpi=200, bbox_inches="tight")
    plt.close()
    print(f'  Saved: {save_path}')


def save_subject_heatmap(subject_summary_df, save_path):
    if subject_summary_df.empty:
        return
    pivot = subject_summary_df.pivot_table(index="subject_category", columns="model_display_name", values="accuracy", aggfunc="mean").sort_index()
    if pivot.empty:
        return
    fig, ax = plt.subplots(figsize=(10, max(4, 0.5 * len(pivot.index))))
    im = ax.imshow(pivot.values, aspect="auto")
    ax.set_xticks(np.arange(len(pivot.columns)))
    ax.set_xticklabels(pivot.columns, rotation=20, ha="right")
    ax.set_yticks(np.arange(len(pivot.index)))
    ax.set_yticklabels(pivot.index)
    ax.set_title("Accuracy by subject/category")
    fig.colorbar(im, ax=ax, label="Accuracy")
    for i in range(pivot.shape[0]):
        for j in range(pivot.shape[1]):
            ax.text(j, i, f"{pivot.iloc[i, j]:.2f}", ha="center", va="center", fontsize=9)
    plt.tight_layout()
    plt.savefig(save_path, dpi=200, bbox_inches="tight")
    plt.close()
    print(f'  Saved: {save_path}')


def save_confusion_matrices(pred_df, figure_dir):
    if pred_df.empty:
        return
    for model_name, grp in pred_df.groupby("model_display_name"):
        matrix = np.zeros((len(LABELS), len(LABELS)), dtype=int)
        for i, gold in enumerate(LABELS):
            for j, pred in enumerate(LABELS):
                matrix[i, j] = int(((grp["gold_answer"] == gold) & (grp["pred_answer"] == pred)).sum())
        fig, ax = plt.subplots(figsize=(5, 4))
        im = ax.imshow(matrix, aspect="auto")
        ax.set_xticks(np.arange(len(LABELS)))
        ax.set_yticks(np.arange(len(LABELS)))
        ax.set_xticklabels(LABELS)
        ax.set_yticklabels(LABELS)
        ax.set_xlabel("Predicted answer")
        ax.set_ylabel("Gold answer")
        ax.set_title(f"Confusion matrix: {model_name}")
        fig.colorbar(im, ax=ax)
        for i in range(len(LABELS)):
            for j in range(len(LABELS)):
                ax.text(j, i, str(matrix[i, j]), ha="center", va="center", fontsize=10)
        plt.tight_layout()
        filename = re.sub(r"[^a-zA-Z0-9_-]+", "_", model_name.lower()) + "_confusion_matrix.png"
        save_path = os.path.join(figure_dir, filename)
        plt.savefig(save_path, dpi=200, bbox_inches="tight")
        plt.close()
        print(f'  Saved: {save_path}')


print(f'\n[STEP 7] Generating figures')
pred_df = ensure_nonempty_csv(PREDICTIONS_PATH)
sbs_df = ensure_nonempty_csv(SUMMARY_BY_SPLIT_PATH)
agg_df = ensure_nonempty_csv(SUMMARY_AGG_PATH)
subj_df = ensure_nonempty_csv(SUBJECT_SUMMARY_PATH)

if not agg_df.empty:
    save_grouped_bar(agg_df, "model_display_name", ["mean_accuracy", "mean_macro_f1"],
                     "Overall model comparison across completed test sets", "Score",
                     os.path.join(FIGURE_DIR, "overall_model_comparison.png"))

if not sbs_df.empty:
    save_accuracy_by_test_set(sbs_df, os.path.join(FIGURE_DIR, "accuracy_by_test_set.png"))

if not pred_df.empty:
    save_latency_boxplot(pred_df, os.path.join(FIGURE_DIR, "latency_distribution.png"))
    save_confusion_matrices(pred_df, FIGURE_DIR)

if not subj_df.empty and subj_df["subject_category"].nunique() > 0:
    save_subject_heatmap(subj_df, os.path.join(FIGURE_DIR, "accuracy_by_subject_heatmap.png"))

print(f'\n{"=" * 60}')
print('DONE — All outputs saved.')
print(f'{"=" * 60}')
print(f'  Predictions CSV:     {PREDICTIONS_PATH}')
print(f'  Summary by split:    {SUMMARY_BY_SPLIT_PATH}')
print(f'  Aggregated summary:  {SUMMARY_AGG_PATH}')
print(f'  Subject summary:     {SUBJECT_SUMMARY_PATH}')
print(f'  Figures dir:         {FIGURE_DIR}')
for p in sorted(glob.glob(os.path.join(FIGURE_DIR, "*.png"))):
    print(f'    - {p}')
