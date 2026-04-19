#!/usr/bin/env python3
"""Convert MCAT DOCX files to JSONL for OpenAI chat fine-tuning.
Extracts embedded images and saves them to images/{category}/."""

import json
import os
import re
import sys
from docx import Document
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(BASE_DIR, "mcat_training.jsonl")
IMAGE_BASE_DIR = os.path.join(BASE_DIR, "images")

SYSTEM_MSG = (
    "You are an MCAT tutor. Solve the question by identifying the key concept, "
    "eliminating incorrect answer choices, and then choosing the best answer."
)

FOLDERS = [
    "BB", "BB2", "CARS_1", "CARS_2", "CARS_d1", "CARS_d2",
    "CP", "Physics",
    "Section_bank_1", "Section_bank_2", "Section_bank_3",
    "Section_bank_4", "Section_bank_5", "Section_bank_6",
]

BOILERPLATE_FRAGMENTS = [
    "Skip to main content",
    "Done Reviewing",
    "Provide Feedback",
    "\u00a9 2026 AAMC",
    "\u00a9 2025 AAMC",
    "\u00a9 2024 AAMC",
    "Contact Us",
    "Web Accessibility",
    "Terms and Conditions",
    "Privacy Statement",
]

SUBJECT_CODE_MAP = {
    "BIO": "BIO", "BCH": "BCH", "GCH": "GCH", "PHY": "PHY",
    "PSY": "PSY", "SOC": "SOC", "ORG": "ORG", "OCH": "OCH",
}

CHOICE_LETTERS = ["A", "B", "C", "D"]


def clean(text):
    return text.replace("\xa0", " ").strip()


def is_boilerplate(text):
    raw = re.sub(r"\[Image: [^\]]+\]", "", text).strip()
    for frag in BOILERPLATE_FRAGMENTS:
        if frag in raw:
            return True
    if re.match(r"^next question\s*\d*$", raw, re.IGNORECASE):
        return True
    return False


def detect_image_refs(text):
    refs = set()
    for m in re.finditer(r"\b(Figure|Fig\.)\s*\d+", text):
        refs.add(m.group(0).strip())
    for m in re.finditer(r"\bTable\s+\d+", text):
        refs.add(m.group(0).strip())
    for m in re.finditer(r"\b(Graph|Diagram|Chart)\s+\d+", text):
        refs.add(m.group(0).strip())
    return refs


def extract_correct_answer_from_solution(text):
    m = re.search(r"correct answer is\s+([A-D])", text, re.IGNORECASE)
    if m:
        return m.group(1).upper()
    m = re.search(r"The solution is\s+([A-D])", text, re.IGNORECASE)
    if m:
        return m.group(1).upper()
    return None


# ---------------------------------------------------------------------------
# Image extraction
# ---------------------------------------------------------------------------

def extract_images_from_doc(doc, folder, file_stem):
    """Extract all images from a DOCX, save them, and map to paragraph indices.

    Returns:
        para_images: dict[int, list[str]]  –  raw paragraph index -> list of
            relative image paths saved to disk.
        all_image_paths: list[str]  –  all saved image paths (relative).
    """
    image_dir = os.path.join(IMAGE_BASE_DIR, folder)

    rid_to_info = {}
    for rid, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            try:
                blob = rel.target_part.blob
                ext = os.path.splitext(rel.target_ref)[1] or ".png"
                rid_to_info[rid] = (blob, ext)
            except Exception:
                pass

    if not rid_to_info:
        return {}, []

    os.makedirs(image_dir, exist_ok=True)

    para_images = {}
    all_image_paths = []
    img_counter = 0
    saved_rids = {}

    for i, para in enumerate(doc.paragraphs):
        blips = para._element.findall(".//" + qn("a:blip"))
        for blip in blips:
            embed = blip.get(qn("r:embed"))
            if not embed or embed not in rid_to_info:
                continue

            if embed in saved_rids:
                rel_path = saved_rids[embed]
            else:
                blob, ext = rid_to_info[embed]
                img_name = f"{file_stem}_{img_counter}{ext}"
                abs_path = os.path.join(image_dir, img_name)
                with open(abs_path, "wb") as f:
                    f.write(blob)
                rel_path = f"images/{folder}/{img_name}"
                saved_rids[embed] = rel_path
                all_image_paths.append(rel_path)
                img_counter += 1

            para_images.setdefault(i, []).append(rel_path)

    return para_images, all_image_paths


# ---------------------------------------------------------------------------
# Main parser
# ---------------------------------------------------------------------------

def parse_docx(filepath, folder):
    """Parse a single DOCX file and return a structured record or None."""
    filename = os.path.basename(filepath)
    file_stem = os.path.splitext(filename)[0]

    try:
        doc = Document(filepath)
    except Exception as e:
        print(f"  SKIP {folder}/{filename}: {e}", file=sys.stderr)
        return None

    # ---- extract & save images, map to paragraphs ----
    para_images, all_image_paths = extract_images_from_doc(doc, folder, file_stem)

    # ---- build raw paragraph list, injecting [Image] markers ----
    raw_paras = []
    for i, p in enumerate(doc.paragraphs):
        txt = clean(p.text)
        if i in para_images:
            markers = " ".join(f"[Image: {ip}]" for ip in para_images[i])
            txt = f"{txt} {markers}".strip() if txt else markers
        raw_paras.append(txt)

    has_image = bool(all_image_paths)
    non_empty = [(i, t) for i, t in enumerate(raw_paras) if t]

    source_title = None
    section_name = None
    question_id_str = None
    passage_header = None
    passage_text_parts = []
    question_stem = None
    choices = {}
    solution_text_parts = []
    correct_answer = None
    user_answer = None
    confidence = None
    content_skills = []
    is_cars_d = folder.startswith("CARS_d")

    # --- Phase 1: locate key markers ---
    question_marker_idx = None
    separator_idx = None
    solution_idx = None
    correct_incorrect_idx = None
    your_answer_idx = None
    correct_answer_idx = None
    confidence_idx = None
    content_skills_idx = None
    option_rationale_idx = None
    item_rationale_idx = None

    for pos, (i, t) in enumerate(non_empty):
        raw_t = re.sub(r"\s*\[Image: [^\]]+\]", "", t).strip()
        if raw_t == "Question":
            question_marker_idx = pos
        elif raw_t == "_":
            separator_idx = pos
        elif raw_t.startswith("Solution:") or raw_t.startswith("The solution is"):
            solution_idx = pos
        elif raw_t in ("Correct", "Incorrect") and solution_idx is not None:
            if correct_incorrect_idx is None:
                correct_incorrect_idx = pos
        elif raw_t == "Your Answer:":
            your_answer_idx = pos
        elif raw_t == "Correct Answer:":
            correct_answer_idx = pos
        elif raw_t == "Confidence Level:":
            confidence_idx = pos
        elif raw_t == "Content & Skills":
            content_skills_idx = pos
        elif raw_t == "Option Rationale:":
            option_rationale_idx = pos
        elif raw_t == "Item Rationale:":
            item_rationale_idx = pos

    if question_marker_idx is None or solution_idx is None:
        print(f"  SKIP {folder}/{filename}: missing Question or Solution marker", file=sys.stderr)
        return None

    # helper: strip image markers for logic checks but keep for content
    def strip_img(t):
        return re.sub(r"\s*\[Image: [^\]]+\]", "", t).strip()

    # --- Phase 2: extract metadata from header ---
    header_end = min(
        question_marker_idx,
        separator_idx if separator_idx is not None else question_marker_idx,
    )

    for pos in range(header_end):
        _, t = non_empty[pos]
        raw_t = strip_img(t)
        if is_boilerplate(raw_t):
            continue
        if raw_t in ("Plan", "Tasks", "Practice", "Resources", "Search"):
            continue
        if raw_t.startswith("Time Spent:"):
            continue
        if re.match(r"^\d+/\d+$", raw_t):
            question_id_str = raw_t.split("/")[0]
            continue
        if source_title is None and pos <= 3:
            source_title = raw_t
            continue
        if section_name is None and source_title is not None and pos <= 6:
            if not re.match(r"^\d+/\d+$", raw_t):
                section_name = raw_t
            continue

    if is_cars_d and section_name is None:
        for pos in range(header_end):
            _, t = non_empty[pos]
            raw_t = strip_img(t)
            if "Diagnostic" in raw_t and "Step" in raw_t:
                section_name = raw_t
                break

    # --- Phase 3: extract passage ---
    passage_start = None
    for pos in range(header_end):
        _, t = non_empty[pos]
        raw_t = strip_img(t)
        if re.match(r"^Passage\s+\d+", raw_t):
            passage_header = raw_t
            passage_start = pos + 1
            break

    if passage_start is not None:
        passage_end = (
            separator_idx
            if separator_idx is not None and separator_idx > passage_start
            else question_marker_idx
        )
        for pos in range(passage_start, passage_end):
            _, t = non_empty[pos]
            raw_t = strip_img(t)
            if raw_t == "_":
                continue
            if is_boilerplate(raw_t):
                continue
            if re.match(r"^Adapted from\b", raw_t) or re.match(r"^Material used in this", raw_t):
                continue
            passage_text_parts.append(t)

    # --- Phase 4: extract question stem ---
    stem_start = question_marker_idx + 1
    stem_end = None
    for pos in range(stem_start, len(non_empty)):
        _, t = non_empty[pos]
        raw_t = strip_img(t)
        if raw_t in CHOICE_LETTERS:
            stem_end = pos
            break
        if raw_t == "_":
            continue

    if stem_end is None:
        print(f"  SKIP {folder}/{filename}: cannot find answer choices", file=sys.stderr)
        return None

    stem_parts = []
    for pos in range(stem_start, stem_end):
        _, t = non_empty[pos]
        raw_t = strip_img(t)
        if raw_t != "_":
            stem_parts.append(t)
    question_stem = " ".join(stem_parts)

    # --- Phase 5: extract choices ---
    choices_start = stem_end
    choices_end = solution_idx

    current_letter = None
    for pos in range(choices_start, choices_end):
        _, t = non_empty[pos]
        raw_t = strip_img(t)
        if raw_t in CHOICE_LETTERS:
            current_letter = raw_t
            img_part = t[len(raw_t):].strip()
            choices[current_letter] = [img_part] if img_part else []
        elif current_letter is not None:
            if raw_t == "Answer choice eliminated":
                continue
            choices[current_letter].append(t)

    for letter in choices:
        choices[letter] = " ".join(choices[letter]).strip()
        if not choices[letter]:
            choices[letter] = "[image]"

    if not choices:
        print(f"  SKIP {folder}/{filename}: no choices extracted", file=sys.stderr)
        return None

    # --- Phase 6: extract correct answer ---
    _, sol_text = non_empty[solution_idx]
    correct_answer = extract_correct_answer_from_solution(strip_img(sol_text))

    # --- Phase 7: extract explanation ---
    explanation_end = (
        correct_incorrect_idx
        if correct_incorrect_idx is not None
        else (your_answer_idx if your_answer_idx is not None else len(non_empty))
    )
    for pos in range(solution_idx, explanation_end):
        _, t = non_empty[pos]
        raw_t = strip_img(t)
        if raw_t.startswith("Solution:") or raw_t.startswith("The solution is"):
            remainder = re.sub(r"^Solution:\s*The correct answer is\s+[A-D]\.\s*", "", raw_t)
            remainder = re.sub(r"^The solution is\s+[A-D]:\s*", "", remainder)
            if remainder.strip():
                solution_text_parts.append(remainder.strip())
        elif raw_t in ("Item Rationale:", "Option Rationale:"):
            continue
        else:
            solution_text_parts.append(raw_t)

    # --- Phase 8: extract user answer, correct answer (verification), confidence ---
    if your_answer_idx is not None and your_answer_idx + 1 < len(non_empty):
        _, ua = non_empty[your_answer_idx + 1]
        if re.match(r"^[A-D]$", strip_img(ua)):
            user_answer = strip_img(ua)

    if correct_answer_idx is not None and correct_answer_idx + 1 < len(non_empty):
        _, ca = non_empty[correct_answer_idx + 1]
        if re.match(r"^[A-D]$", strip_img(ca)):
            if correct_answer is None:
                correct_answer = strip_img(ca)

    if confidence_idx is not None and confidence_idx + 1 < len(non_empty):
        _, conf = non_empty[confidence_idx + 1]
        raw_conf = strip_img(conf)
        if not raw_conf.startswith("Content") and raw_conf != "Not Selected":
            confidence = raw_conf
        else:
            confidence = None

    # --- Phase 9: extract content & skills ---
    if content_skills_idx is not None:
        for pos in range(content_skills_idx + 1, len(non_empty)):
            _, t = non_empty[pos]
            raw_t = strip_img(t)
            if is_boilerplate(raw_t):
                break
            if re.match(r"^next question", raw_t, re.IGNORECASE):
                break
            if re.match(r"^[A-D]\s+selected$", raw_t, re.IGNORECASE):
                break
            tag = re.sub(r"^Skill:\s*", "", raw_t).strip()
            if tag:
                content_skills.append(tag)

    # --- Phase 10: determine subject_code ---
    subject_code = None
    for tag in reversed(content_skills):
        upper_tag = tag.upper()
        if upper_tag in SUBJECT_CODE_MAP:
            subject_code = SUBJECT_CODE_MAP[upper_tag]
            break
    if subject_code is None and folder.startswith("CARS"):
        subject_code = "CARS"
    if subject_code is None and folder == "Physics":
        subject_code = "PHY"
    if subject_code is None and folder == "CP":
        subject_code = "GCH"

    # --- Phase 11: image notes ---
    all_text = " ".join(
        [question_stem or ""] + list(choices.values()) + solution_text_parts + passage_text_parts
    )
    image_refs = detect_image_refs(all_text)
    image_note = None
    if has_image and image_refs:
        image_note = f"Contains {', '.join(sorted(image_refs))}; {len(all_image_paths)} image(s) extracted"
    elif has_image:
        image_note = f"{len(all_image_paths)} embedded image(s) extracted"
    elif image_refs:
        has_image = True
        image_note = f"References {', '.join(sorted(image_refs))} but no embedded images found"

    # --- Phase 12: build elimination reasoning ---
    elimination = build_elimination(
        choices, correct_answer, solution_text_parts,
        is_cars_d, non_empty, option_rationale_idx, correct_incorrect_idx,
    )

    # --- Phase 13: build concept ---
    concept = build_concept(subject_code, content_skills, question_stem, solution_text_parts)

    # --- Phase 14: build final explanation / "Why" ---
    why_text = build_why(correct_answer, solution_text_parts, elimination, is_cars_d)

    # --- Phase 15: format messages ---
    user_content = format_user_message(question_stem, choices, passage_text_parts)
    assistant_content = format_assistant_message(concept, elimination, correct_answer, why_text)

    record = {
        "messages": [
            {"role": "system", "content": SYSTEM_MSG},
            {"role": "user", "content": user_content},
            {"role": "assistant", "content": assistant_content},
        ],
        "metadata": {
            "source_file": filename,
            "category": folder,
            "source_title": source_title,
            "section": section_name,
            "subject_code": subject_code,
            "question_id": question_id_str,
            "correct_answer": correct_answer,
            "user_answer": user_answer,
            "confidence": confidence,
            "content_skills": content_skills if content_skills else None,
            "has_image": has_image,
            "images": all_image_paths if all_image_paths else None,
            "image_note": image_note,
        },
    }

    return record


# ---------------------------------------------------------------------------
# Reasoning builders
# ---------------------------------------------------------------------------

def build_elimination(choices, correct_answer, solution_parts, is_cars_d,
                      non_empty, option_rationale_idx, correct_incorrect_idx):
    """Build per-choice elimination reasoning."""
    elim = {}

    def strip_img(t):
        return re.sub(r"\s*\[Image: [^\]]+\]", "", t).strip()

    if is_cars_d and option_rationale_idx is not None:
        end_idx = correct_incorrect_idx if correct_incorrect_idx is not None else len(non_empty)
        current_opt = None
        current_text = []
        for pos in range(option_rationale_idx + 1, end_idx):
            _, t = non_empty[pos]
            raw_t = strip_img(t)
            m = re.match(r"^Option\s+([A-D]):\s*(.*)", raw_t)
            if m:
                if current_opt:
                    elim[current_opt] = " ".join(current_text).strip()
                current_opt = m.group(1)
                current_text = [m.group(2)] if m.group(2) else []
            elif current_opt:
                current_text.append(raw_t)
        if current_opt:
            elim[current_opt] = " ".join(current_text).strip()

        for letter in elim:
            text = elim[letter]
            is_correct = letter == correct_answer
            text = re.sub(r"^(Correct|Incorrect)\.\s*", "", text)
            choice_text = re.sub(r"\[Image: [^\]]+\]", "", choices.get(letter, "")).strip()
            if choice_text and text.startswith(choice_text):
                text = text[len(choice_text):].strip()
                if text and text[0] == ".":
                    text = text[1:].strip()
            text = re.sub(r"^(Correct|Incorrect)\.\s*", "", text)
            prefix = "Correct. " if is_correct else "Incorrect. "
            elim[letter] = prefix + text.strip() if text.strip() else prefix.strip()

    if not elim:
        full_explanation = " ".join(solution_parts)
        per_choice_paragraphs = try_map_paragraphs_to_choices(solution_parts, choices)
        if per_choice_paragraphs:
            elim = per_choice_paragraphs
        else:
            elim = extract_choice_reasoning_from_text(full_explanation, choices, correct_answer)

    for letter in choices:
        if letter not in elim or not elim[letter]:
            if letter == correct_answer:
                elim[letter] = "Correct. See explanation below."
            else:
                elim[letter] = "Incorrect based on the information provided."

    return elim


def try_map_paragraphs_to_choices(solution_parts, choices):
    if len(solution_parts) < len(choices):
        return None
    if len(solution_parts) == len(choices):
        result = {}
        letters = sorted(choices.keys())
        for letter, para in zip(letters, solution_parts):
            text = re.sub(r"^(Correct|Incorrect)\.\s*", "", para.strip())
            result[letter] = text
        return result
    if len(solution_parts) == len(choices) + 1:
        result = {}
        letters = sorted(choices.keys())
        for idx, letter in enumerate(letters):
            para = re.sub(r"^(Correct|Incorrect)\.\s*", "", solution_parts[idx + 1].strip())
            result[letter] = para
        return result
    return None


def extract_choice_reasoning_from_text(text, choices, correct_answer):
    elim = {}
    sentences = re.split(r"(?<=[.!?])\s+", text)

    assigned_to_wrong = set()
    for letter in choices:
        if letter == correct_answer:
            continue
        patterns = [
            rf"\({letter}\)",
            rf"choice\s+{letter}\b",
            rf"option\s+{letter}\b",
            rf"answer\s+{letter}\b",
        ]
        relevant = []
        for idx, sent in enumerate(sentences):
            for pat in patterns:
                if re.search(pat, sent, re.IGNORECASE):
                    relevant.append(sent)
                    assigned_to_wrong.add(idx)
                    break
        if relevant:
            combined = " ".join(relevant)
            combined = re.sub(r"^(Correct|Incorrect)\.\s*", "", combined)
            if len(combined) > 400:
                combined = combined[:397] + "..."
            elim[letter] = combined

    if correct_answer:
        correct_sentences = []
        for idx, sent in enumerate(sentences):
            if idx not in assigned_to_wrong:
                cleaned = sent.strip()
                if re.match(r"^Thus,?\s+[A-D]\s+is\s+(the\s+)?(best|correct)\s+answer", cleaned, re.IGNORECASE):
                    continue
                if re.match(r"^(Choice|Option)\s+[A-D]\s+is\s+(the\s+)?(best|correct)", cleaned, re.IGNORECASE):
                    continue
                if cleaned:
                    correct_sentences.append(cleaned)
        if correct_sentences:
            combined = " ".join(correct_sentences)
            if len(combined) > 500:
                combined = combined[:497] + "..."
            elim[correct_answer] = "Correct. " + combined

    return elim


def build_concept(subject_code, content_skills, question_stem, solution_parts):
    code_name = {
        "BIO": "Biology",
        "BCH": "Biochemistry",
        "GCH": "General Chemistry",
        "PHY": "Physics",
        "PSY": "Psychology",
        "SOC": "Sociology",
        "ORG": "Organic Chemistry",
        "OCH": "Organic Chemistry",
        "CARS": "Critical Analysis and Reasoning Skills",
    }.get(subject_code, "MCAT")

    full_text = " ".join(solution_parts)
    m = re.match(r"^(This is an?\s+.+?question\s+that\s+.+?\.)", full_text)
    if not m:
        m = re.match(r"^(This question\s+.+?\.)", full_text)
    if m:
        first_sent = m.group(1)
        first_sent = re.sub(
            r"^This is an?\s+\w+\s+question\s+that\s+falls under the content category\s*", "", first_sent
        )
        first_sent = re.sub(r'^"(.+?)"\.?\s*', r"\1. ", first_sent)
        if first_sent.strip():
            return f"{code_name} \u2014 {first_sent.strip()}"

    if question_stem:
        stem_clean = re.sub(r"\[Image: [^\]]+\]", "", question_stem).strip()
        stem_short = stem_clean[:120]
        if len(stem_clean) > 120:
            stem_short = stem_short.rsplit(" ", 1)[0] + "..."
        return f"{code_name} \u2014 reasoning about: {stem_short}"

    return f"{code_name} concept application."


def build_why(correct_answer, solution_parts, elimination, is_cars_d=False):
    if is_cars_d:
        item_rationale_parts = []
        in_item = False
        for part in solution_parts:
            if re.match(r"^This is a\s+.+\s+question\s+because", part):
                item_rationale_parts.append(part)
                in_item = True
                continue
            if re.match(r"^Option\s+[A-D]:", part):
                in_item = False
            if in_item:
                item_rationale_parts.append(part)

        correct_reason = ""
        if correct_answer and correct_answer in elimination:
            correct_reason = re.sub(r"^(Correct|Incorrect)\.\s*", "", elimination[correct_answer])

        parts = []
        if item_rationale_parts:
            parts.append(" ".join(item_rationale_parts))
        if correct_reason:
            parts.append(correct_reason)
        full = " ".join(parts).strip()
        if not full:
            full = " ".join(solution_parts).strip()
    else:
        full = " ".join(solution_parts).strip()

    full = re.sub(
        r"^This is an?\s+\w+\s+question\s+that\s+falls under\s+the\s+content\s+category\s*\"[^\"]*\"\.\s*",
        "", full,
    )
    full = re.sub(r"^The answer to this question is [A-D] because\s+", "", full)
    full = re.sub(r"\s*It is a [\w\s]+question because it requires.*$", "", full)
    full = full.strip()
    full = re.sub(r"^Thus,?\s+[A-D]\s+is\s+(the\s+)?(best|correct)\s+answer\.?\s*$", "", full).strip()

    if not full and correct_answer and correct_answer in elimination:
        text = elimination[correct_answer]
        text = re.sub(r"^(Correct|Incorrect)\.\s*", "", text)
        full = text

    if len(full) > 1500:
        full = full[:1497] + "..."

    return full if full else "See elimination reasoning above."


# ---------------------------------------------------------------------------
# Message formatters
# ---------------------------------------------------------------------------

def format_user_message(question_stem, choices, passage_parts):
    parts = []
    if passage_parts:
        parts.append("Passage:")
        parts.append("\n".join(passage_parts))
        parts.append("")
    parts.append("Question:")
    parts.append(question_stem)
    parts.append("")
    parts.append("Choices:")
    for letter in sorted(choices.keys()):
        parts.append(f"{letter}. {choices[letter]}")
    return "\n".join(parts)


def format_assistant_message(concept, elimination, correct_answer, why_text):
    parts = [
        "Concept:",
        concept,
        "",
        "Elimination:",
    ]
    for letter in sorted(elimination.keys()):
        parts.append(f"{letter}. {elimination[letter]}")
    parts += [
        "",
        "Answer:",
        correct_answer if correct_answer else "Unknown",
        "",
        "Why:",
        why_text,
    ]
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Validation
# ---------------------------------------------------------------------------

def validate_record(record, folder, filename):
    errors = []
    msgs = record.get("messages", [])
    if len(msgs) != 3:
        errors.append(f"Expected 3 messages, got {len(msgs)}")
    if msgs[0]["role"] != "system":
        errors.append("msg[0] must be system")
    if msgs[1]["role"] != "user":
        errors.append("msg[1] must be user")
    if msgs[2]["role"] != "assistant":
        errors.append("msg[2] must be assistant")

    meta = record.get("metadata", {})
    ca = meta.get("correct_answer")
    if ca is None:
        errors.append("Missing correct_answer")

    assistant_text = msgs[2]["content"]
    answer_match = re.search(r"^Answer:\n(.+)$", assistant_text, re.MULTILINE)
    if answer_match:
        if ca and answer_match.group(1).strip() != ca:
            errors.append(f"Answer mismatch: metadata={ca}, text={answer_match.group(1).strip()}")

    if "Choices:" not in msgs[1]["content"]:
        errors.append("User message missing Choices section")

    if errors:
        for e in errors:
            print(f"  WARN {folder}/{filename}: {e}", file=sys.stderr)
        return False
    return True


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    os.makedirs(IMAGE_BASE_DIR, exist_ok=True)

    total = 0
    written = 0
    skipped = 0
    warnings = 0
    images_saved = 0

    with open(OUTPUT_FILE, "w", encoding="utf-8") as out:
        for folder in FOLDERS:
            folder_path = os.path.join(BASE_DIR, folder)
            if not os.path.isdir(folder_path):
                print(f"Folder not found: {folder}", file=sys.stderr)
                continue

            files = sorted(
                [f for f in os.listdir(folder_path)
                 if f.endswith(".docx") and not f.startswith("~$")],
                key=lambda x: int(re.search(r"\d+", x).group()) if re.search(r"\d+", x) else 0,
            )

            print(f"Processing {folder}: {len(files)} files...", file=sys.stderr)

            for fname in files:
                total += 1
                fpath = os.path.join(folder_path, fname)
                record = parse_docx(fpath, folder)

                if record is None:
                    skipped += 1
                    continue

                valid = validate_record(record, folder, fname)
                if not valid:
                    warnings += 1

                img_list = record["metadata"].get("images")
                if img_list:
                    images_saved += len(img_list)

                line = json.dumps(record, ensure_ascii=False)
                out.write(line + "\n")
                written += 1

    print(
        f"\nDone. Total: {total}, Written: {written}, Skipped: {skipped}, "
        f"Warnings: {warnings}, Images saved: {images_saved}",
        file=sys.stderr,
    )
    print(f"Output: {OUTPUT_FILE}", file=sys.stderr)
    print(f"Images: {IMAGE_BASE_DIR}/", file=sys.stderr)


if __name__ == "__main__":
    main()
